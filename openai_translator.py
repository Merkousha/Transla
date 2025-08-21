import os
import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from dotenv import load_dotenv
import asyncio
from openai import OpenAI
import re

# Load environment variables
load_dotenv()

class OpenAIBookTranslator:
    def __init__(self, api_key=None, model="gpt-4o", base_url="https://api.avalai.ir/v1"):
        """
        Initialize the book translator with OpenAI-compatible service
        
        Args:
            api_key (str): OpenAI API key
            model (str): Model name (e.g., "gpt-4o", "gpt-3.5-turbo")
            base_url (str): Base URL for the service (default: Avalai)
        """
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("API key is required. Set OPENAI_API_KEY environment variable or pass it to constructor.")
        
        self.model = model
        self.base_url = base_url
        
        # Configure OpenAI client
        self.client = OpenAI(
            api_key=self.api_key,
            base_url=self.base_url
        )
        
        print(f"Initialized OpenAI client with model: {self.model}")
        print(f"Using endpoint: {self.base_url}")
        
    def extract_text_from_pdf_with_overlap(self, pdf_path, overlap_paragraphs=1):
        """
        Extract text from PDF file with overlapping paragraphs between pages
        """
        print(f"Reading PDF file: {pdf_path}")
        pages_text = []
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    print(f"Processing page {page_num}/{total_pages}")
                    page_text = page.extract_text()
                    if page_text.strip():
                        pages_text.append(page_text.strip())
                        
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return None
        
        # Process pages with overlapping paragraphs
        processed_text = self._merge_pages_with_overlap(pages_text, overlap_paragraphs)
        return processed_text
    
    def _merge_pages_with_overlap(self, pages_text, overlap_paragraphs=1):
        """
        Merge pages with overlapping paragraphs to maintain context
        """
        print("Merging pages with overlapping paragraphs...")
        
        if not pages_text:
            return ""
        
        merged_text = ""
        previous_paragraphs = []
        
        for i, page_text in enumerate(pages_text):
            # Split page into paragraphs
            paragraphs = self._split_into_paragraphs(page_text)
            
            if not paragraphs:
                continue
            
            # Add overlapping paragraphs from previous page
            if previous_paragraphs:
                overlap_text = "\n\n".join(previous_paragraphs[-overlap_paragraphs:])
                if overlap_text.strip():
                    merged_text += f"[CONTEXT FROM PREVIOUS PAGE: {overlap_text}]\n\n"
            
            # Add current page content
            current_page_text = "\n\n".join(paragraphs)
            merged_text += current_page_text + "\n\n"
            
            # Store last few paragraphs for next page overlap
            previous_paragraphs = paragraphs[-overlap_paragraphs:] if len(paragraphs) >= overlap_paragraphs else paragraphs
        
        return merged_text
    
    def _split_into_paragraphs(self, text):
        """
        Split text into paragraphs while preserving structure
        """
        # Split by double newlines (common paragraph separator)
        paragraphs = re.split(r'\n\s*\n', text)
        
        # Clean up paragraphs
        cleaned_paragraphs = []
        for para in paragraphs:
            para = para.strip()
            if para and len(para) > 10:  # Only keep substantial paragraphs
                cleaned_paragraphs.append(para)
        
        return cleaned_paragraphs
    
    def extract_text_from_pdf(self, pdf_path):
        """
        Legacy method - now calls the new method with overlap
        """
        return self.extract_text_from_pdf_with_overlap(pdf_path, overlap_paragraphs=1)
    
    def split_text_into_chunks(self, text, chunk_size=2000, chunk_overlap=200):
        """
        Split text into manageable chunks for translation (simple implementation)
        """
        print("Splitting text into chunks...")
        
        # Simple text splitting by sentences and paragraphs
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Add period back to sentence
            sentence += "."
            
            if len(current_chunk) + len(sentence) < chunk_size:
                current_chunk += " " + sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        print(f"Created {len(chunks)} chunks for translation")
        return chunks
    
    async def translate_chunk(self, text, source_lang="English", target_lang="Persian"):
        """
        Translate a text chunk using OpenAI-compatible service
        """
        try:
            system_prompt = f"""You are a professional translator specializing in software and technical documentation. Translate the following text from {source_lang} to {target_lang}.

IMPORTANT TRANSLATION RULES:
1. Maintain the original formatting, structure, and meaning
2. DO NOT translate technical software terms - keep them in English:
   - Programming concepts: Run, Build, Deploy, Debug, Compile, Execute, Test, Refactor, Optimize
   - Software patterns: Domain Driven Design, MVC, MVP, MVVM, Repository Pattern, Factory Pattern
   - Development terms: Sprint, Backlog, User Story, Epic, Bug, Feature, Hotfix, Release
   - Technical terms: API, SDK, Framework, Library, Module, Package, Dependency, Version
   - Code-related: Function, Method, Class, Object, Variable, Parameter, Return, Import, Export
   - Tools: Git, Docker, Kubernetes, Jenkins, Jira, VS Code, IntelliJ
   - Platforms: AWS, Azure, Google Cloud, Heroku, DigitalOcean, GitHub, GitLab, Bitbucket
   - Languages: Python, JavaScript, Java, C#, TypeScript, React, Angular, Vue, Node.js
   - Databases: MySQL, PostgreSQL, MongoDB, Redis, SQLite, Oracle, SQL Server
   - Protocols: HTTP, HTTPS, REST, GraphQL, WebSocket, TCP, UDP, SSH, FTP
   - File formats: JSON, XML, CSV, YAML, Markdown, HTML, CSS, SVG, PNG, JPG
   - Version control: Commit, Push, Pull, Merge, Branch, Fork, Clone, Repository
   - DevOps: CI/CD, Pipeline, Container, Microservices, Monolith, Serverless
   - Testing: Unit Test, Integration Test, E2E Test, Mock, Stub, Fixture, Assertion
   - Architecture: Monorepo, Microfrontend, Service Mesh, Event Sourcing, CQRS
   - Security: Authentication, Authorization, Encryption, JWT, OAuth, SSL, TLS
   - Performance: Caching, Load Balancing, Scaling, Optimization, Benchmarking
   - Monitoring: Logging, Metrics, Alerting, Tracing, Profiling, Health Check
   - Deployment: Blue-Green, Canary, Rolling Update, Rollback, Zero-Downtime

3. If there are technical terms not in the above list, use your judgment but prefer keeping them in English if they are commonly used in software development
4. If the text contains numbers, dates, or proper nouns, keep them as is
5. If you see [CONTEXT FROM PREVIOUS PAGE: ...], translate that context as well to maintain continuity
6. Only provide the translation, no explanations or additional text

Examples of correct translation:
- "Run the application" → "اپلیکیشن را Run کنید"
- "Build the project" → "پروژه را Build کنید"
- "Domain Driven Design principles" → "اصول Domain Driven Design"
- "Deploy to production" → "Deploy کردن به production"
- "Create a new branch" → "ایجاد یک branch جدید"
- "Merge the changes" → "Merge کردن تغییرات"
- "Debug the issue" → "Debug کردن مشکل"
- "Test the functionality" → "Test کردن قابلیت"
- "API endpoint" → "API endpoint"
- "Database connection" → "اتصال Database"
- "Git repository" → "Git repository"
- "Docker container" → "Docker container"
- "REST API" → "REST API"
- "JSON response" → "پاسخ JSON"
- "HTTP request" → "درخواست HTTP"
- "SSL certificate" → "گواهی SSL"
- "Unit test" → "Unit test"
- "Code review" → "بازبینی Code"
- "Pull request" → "Pull request"
- "Release notes" → "یادداشت‌های Release"

Text to translate:
{text}

Translation:"""
            
            user_prompt = f"Text to translate:\n{text}"
            
            # Use OpenAI-compatible service
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"Translation error: {e}")
            return f"[Translation Error: {e}]"
    
    def create_word_document(self, original_chunks, translated_chunks, output_path):
        """
        Create a Word document with original and translated text
        """
        print(f"Creating Word document: {output_path}")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Book Translation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Original Language: English")
        doc.add_paragraph(f"Target Language: Persian")
        doc.add_paragraph(f"Total Sections: {len(original_chunks)}")
        doc.add_paragraph(f"Translation Engine: OpenAI-Compatible Service")
        doc.add_paragraph(f"Model: {self.model}")
        doc.add_paragraph(f"Endpoint: {self.base_url}")
        doc.add_paragraph(f"Page Overlap: Enabled (for better continuity)")
        doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")  # Empty line
        
        # Add content
        for i, (original, translated) in enumerate(zip(original_chunks, translated_chunks), 1):
            # Section header
            section_header = doc.add_heading(f'Section {i}', level=1)
            
            # Original text
            doc.add_heading('Original Text:', level=2)
            original_para = doc.add_paragraph(original)
            original_para.style = 'Normal'
            
            # Translated text
            doc.add_heading('Translated Text:', level=2)
            translated_para = doc.add_paragraph(translated)
            translated_para.style = 'Normal'
            
            # Add separator
            doc.add_paragraph("=" * 80)
            doc.add_paragraph("")  # Empty line
        
        # Save document
        doc.save(output_path)
        print(f"Word document saved successfully: {output_path}")
    
    async def translate_book_async(self, pdf_path, output_path, source_lang="English", target_lang="Persian", overlap_paragraphs=1):
        """
        Main async method to translate a book from PDF to Word document
        """
        print("Starting book translation process with OpenAI-compatible service...")
        print(f"Using page overlap: {overlap_paragraphs} paragraph(s)")
        
        # Extract text from PDF with overlap
        text = self.extract_text_from_pdf_with_overlap(pdf_path, overlap_paragraphs)
        if not text:
            print("Failed to extract text from PDF")
            return False
        
        # Split text into chunks
        chunks = self.split_text_into_chunks(text)
        
        # Translate chunks
        translated_chunks = []
        total_chunks = len(chunks)
        
        print(f"Starting translation of {total_chunks} chunks...")
        
        for i, chunk in enumerate(chunks, 1):
            print(f"Translating chunk {i}/{total_chunks}")
            print(f"Chunk preview: {chunk[:100]}...")
            
            translated = await self.translate_chunk(chunk, source_lang, target_lang)
            translated_chunks.append(translated)
            
            # Add delay to avoid rate limiting
            if i < total_chunks:
                await asyncio.sleep(1)
        
        # Create Word document
        self.create_word_document(chunks, translated_chunks, output_path)
        
        print("Translation completed successfully!")
        return True
    
    def translate_book(self, pdf_path, output_path, source_lang="English", target_lang="Persian", overlap_paragraphs=1):
        """
        Synchronous wrapper for the async translation method
        """
        return asyncio.run(self.translate_book_async(pdf_path, output_path, source_lang, target_lang, overlap_paragraphs))

def main():
    """
    Main function to run the OpenAI book translator
    """
    # Configuration
    pdf_path = "book.pdf"  # Your PDF file
    output_path = "translated_book_openai.docx"  # Output Word document
    overlap_paragraphs = 1  # Number of paragraphs to overlap between pages
    
    try:
        # Initialize translator
        translator = OpenAIBookTranslator()
        
        # Translate the book
        success = translator.translate_book(pdf_path, output_path, overlap_paragraphs=overlap_paragraphs)
        
        if success:
            print(f"\n✅ Translation completed!")
            print(f"📄 Original PDF: {pdf_path}")
            print(f"📝 Translated Word document: {output_path}")
            print(f"🔄 Page overlap: {overlap_paragraphs} paragraph(s)")
            print(f"🌐 Service: OpenAI-Compatible (https://api.avalai.ir/v1)")
        else:
            print("❌ Translation failed!")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        print("\nMake sure you have:")
        print("1. Set your API key in .env file: OPENAI_API_KEY=your-api-key")
        print("2. Installed all required packages: pip install -r requirements.txt")
        print("3. The PDF file exists in the current directory")
        print("4. You have access to the OpenAI-compatible service")

if __name__ == "__main__":
    main()
